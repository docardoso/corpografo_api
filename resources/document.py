from sqlalchemy.orm import joinedload
import io
import bs4
import pypdf
import docx
import pathlib
import base64
from nltk.util import guess_encoding

import collections as cl
import re
from flask_smorest import Blueprint, abort
from flask.views import MethodView
from flask_jwt_extended import get_jwt_identity

from db import db
from models import Document, User, DocumentsUsers
from schemas import DocumentSchema, PlainDocumentSchema, LightweightDocumentSchema
from util import login_required, get_current_user, get_user_document

from sqlalchemy.exc import IntegrityError

blp = Blueprint("Documents", __name__, description="Operations on documents")

@blp.route("/document/<int:document_id>")
class DocumentPicker(MethodView):
    @login_required()
    @blp.response(200, DocumentSchema)
    def get(self, document_id):
        return get_user_document(document_id)

    #def delete(self, document_id):
    #    db.session.delete(Document.get_or_404(document_id))
    #    db.session.commit()
    #    return {'message': 'Document deleted'}

    @login_required()
    @blp.arguments(DocumentSchema)
    @blp.response(200, DocumentSchema)
    def put(self, document_data, document_id):
        document = db.session.query(DocumentsUsers).get_or_404({
            'document_id':document_id,
            'user_id':get_jwt_identity(),
        }).document

        for k in document_data:
            setattr(document, k, document_data[k])

        db.session.add(document)
        db.session.commit()
        return document

content_extractors = {
    '.txt':  lambda raw_content: guess_encoding(raw_content)[0],
    '.pdf':  lambda raw_content: '\n'.join(i.extract_text() for i in pypdf.PdfReader(io.BytesIO(raw_content)).pages),
    '.html': lambda raw_content: bs4.BeautifulSoup(guess_encoding(raw_content)[0]).get_text(),
    '.htm':  lambda raw_content: bs4.BeautifulSoup(guess_encoding(raw_content)[0]).get_text(),
    '.docx': lambda raw_content: '\n'.join(i.text for i in docx.Document(io.BytesIO(raw_content)).paragraphs),
}

@blp.route("/document")
class DocumentMethodView(MethodView):
    @login_required()
    @blp.arguments(PlainDocumentSchema)
    @blp.response(201, PlainDocumentSchema)
    def post(self, document_data):
        document = Document(**document_data)

        extension = pathlib.Path(document.name).suffix
        try:
            document.content = content_extractors[extension](base64.b64decode(document.input_file.encode('utf-8')))
        except KeyError:
            document.content = ''


        document.users.append(get_current_user())
        db.session.add(document)
        try:
            db.session.commit()
        except IntegrityError as e:
            #db.session.rollback()
            abort(400, message=e._message()) #  'The name of each document must be unique.')  # 

        return document

    @login_required()
    @blp.response(200, LightweightDocumentSchema(many=True))
    def get(self):
        return get_current_user(joinedload(User.document_relations).joinedload(DocumentsUsers.document)).documents

@blp.route('/document/<int:document_id>/user/<user_email>')
class DocumentUserEmail(MethodView):
    @login_required()
    def post(self, document_id, user_email):
        get_user_document(document_id)
        user_id = db.session.query(User).filter_by(email=user_email).first().id
        db.session.add(DocumentsUsers(document_id=document_id, user_id=user_id))
        db.session.commit()

        return {'message': 'Document shared with user'}

@blp.route('/document/<int:document_id>/user/<int:user_id>')
class DocumentUserId(MethodView):
    @login_required()
    def delete(self, document_id, user_id):
        get_user_document(document_id)
        db.session.delete(
            db.session.query(DocumentsUsers).get({
                'document_id': document_id,
                'user_id': user_id,
            })
        )
        db.session.commit()

        return {'message': 'Document unshared with user'}

@blp.route('/dictionary/<int:document_id>')
class Dictionary(MethodView):
    @login_required()
    def get(self, document_id):
        document = get_user_document(document_id)
        return [document.name, cl.Counter(re.findall(r'\w+', document.content.lower()))]

@blp.route('/phrasing/<int:document_id>')
class Phrasing(MethodView):
    @login_required()
    def get(self, document_id):
        document = get_user_document(document_id)
        return [document.name, re.findall(r'.+', document.content)]